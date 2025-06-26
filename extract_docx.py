import os
import re
from datetime import datetime
from docx import Document
from docx.text.run import Run
from docx.text.hyperlink import Hyperlink
from config import DYNAMIC_FASSUNGSLINK_PREFIX

DOC_DIR = 'assets/docx_files'
SKIP_PARA_TEXT_STARTSWITH = ('Buch', '(Stand: ', 'Stand: ', 'Kommentar *', 'Kommentar zu ')
SKIP_PARA_TEXT_STARTSWITH_AUTHOR = ('Stefan Abel:', 'Miriam Strieder:', 'Agata Mazurek:', 'Ann-Kathrin Deininger:')


def extract_docs(doc_dir=DOC_DIR):
    docx_files = [f for f in os.listdir(doc_dir) if f.endswith('.docx') and not f.startswith('~$') or print(f"{f} skipped")]

    docx_files.sort(key=lambda x: x.upper())

    documents = []
    for docx_file in docx_files:
        docx_path = os.path.join(doc_dir, docx_file)
        doc = Doc(docx_path)
        documents.append(doc)
    return documents


class Doc:
    def __init__(self, file):
        self.file_name = file.split('/')[-1]
        self.fassung = 'a' if self.file_name.upper().startswith('DMGT') else self.file_name[0]
        self._last_mod_date = os.path.getmtime(file)
        self.last_mod_date = datetime.fromtimestamp(os.path.getmtime(file)).strftime('%Y.%m.%d')

        self._document = Document(file)
        self._content, self._zitierte_literatur = self.get_content()

        self.paragraphs = self.get_paragraphs(self._content)
        self.literature_paragraphs = self.get_paragraphs(self._zitierte_literatur)

    def get_clean_paragraphs(self):
        return [p for p in self._document.paragraphs
                if p.text and not p.text.startswith(SKIP_PARA_TEXT_STARTSWITH)]

    # get all paragraphs belonging to "zitierte Literatur"
    def get_content(self):
        clean_paragraphs = self.get_clean_paragraphs()
        content = []
        zitierte_literatur = []
        is_literatur = False

        for para in clean_paragraphs:
            if para.text.startswith('Zitierte Literatur:'):
                is_literatur = True
            if is_literatur:
                zitierte_literatur.append(para)
            else:
                if not para.text.startswith(SKIP_PARA_TEXT_STARTSWITH_AUTHOR):  # additional filter, as Stefan Abel might be also a cited Author'
                    content.append(para)
        return content, zitierte_literatur

    def get_paragraphs(self, content):
        c_paragraphs = []
        for idx, p in enumerate(content):
            c_paragraphs.append(Paragraph(p, idx, self.fassung))
        return c_paragraphs


class Paragraph:
    def __init__(self, docx_p, idx, fassung):
        self._docx_paragraph = docx_p
        self.idx = idx
        self.fassung = fassung
        self.text = docx_p.text
        self.comment_elements = self.get_comment_elements(docx_p)
        self.is_title = self.get_is_title()
        self.is_list_item = docx_p.style.style_id == 'ListParagraph'

    def get_comment_elements(self, paragraph):
        comment_parts = []
        idx = 0

        for item in paragraph.iter_inner_content():
            if isinstance(item, Run):
                comment_parts.append(
                    ParagraphElement(idx, item)
                )
                idx += 1
            elif isinstance(item, Hyperlink):
                h = PararaphHyperlinkElement(idx, item)
                for run in item.runs:
                    h.elements.append(
                        ParagraphElement(idx, run, item.url)
                    )
                    idx += 1

                comment_parts.append(h)
        return comment_parts


    # check if any paragraph is a title. For being a title, the paragraph must be bold and start with a number
    # of 1 to 3 digits, followed by a dot and a number of 1 to 2 digits. The number of the second part must be
    # less than 31
    def get_is_title(self):
        match = re.match(r'^(\d{1,3})\.(\d{1,2})', self.text.split(' ')[0])
        if match and self.comment_elements[0].style.bold:
            first_part = int(match.group(1))
            second_part = int(match.group(2))
            return second_part < 31
        return False


class ParagraphElement:
    def __init__(self, idx, run, href=None):
        self.idx = idx
        self.text = run.text
        self.style = StyleAdapter(run)
        self.h_ref = href
        self.has_versnumber_for_citing = self.get_has_versnumber_for_citing()

    # add dynamic hyperlinks to elements containing no hyperlink but a verse number of format 113.2 and are bold
    def get_has_versnumber_for_citing(self):
        if self.h_ref:
            return False

        match = re.match(r'^(\d{1,3})\.(\d{1,2})', self.text.strip())
        if match: #and int(match.group(0).split('.')[0]) < 828 and int(match.group(0).split('.')[1]) < 31:
            #return f"{DYNAMIC_FASSUNGSLINK_PREFIX}/{match.group(1)}"
            return True



class PararaphHyperlinkElement:
    def __init__(self, idx, hyperlink):
        self.idx = idx
        self.text = hyperlink.text
        self.h_ref = hyperlink.url
        self.elements = []


# Adapter class to extract the relevant styles
class StyleAdapter:
    def __init__(self, run):
        self._run = run
        self.bold = True if run.bold else False
        self.italic = True if run.italic else False
        self.small_caps = run.font.small_caps if run.font.small_caps else False
        self.classes = []

        if self.bold:
            self.classes.append('fk-bold')
        if self.italic:
            self.classes.append('fk-italic')
        if self.small_caps:
            self.classes.append('fk-small-caps')

